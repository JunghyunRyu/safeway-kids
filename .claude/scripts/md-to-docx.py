#!/usr/bin/env python3
"""
Markdown → DOCX converter tailored for SafeWay Kids regulatory artifacts.

Usage:
    python md-to-docx.py <input.md> <output.docx> [--title "<title>"]

Supported markdown elements:
- Headings (# through ####)
- Paragraphs
- Bullet lists (- ...)
- Numbered lists (1. ...)
- Block quotes (> ...)
- Tables (| col | col |)
- Bold (**text**) and inline code (`text`)
- Horizontal rules (---)
- Fenced code blocks (``` ... ```)
- YAML frontmatter (--- ... ---) — skipped
"""
import sys
import re
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


KOREAN_FONT = "맑은 고딕"
LATIN_FONT = "Calibri"


def set_cell_font(cell, font_name_latin=LATIN_FONT, font_name_ko=KOREAN_FONT, size=10, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name_latin
            run.font.size = Pt(size)
            run.bold = bold
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), font_name_ko)
            rFonts.set(qn("w:ascii"), font_name_latin)
            rFonts.set(qn("w:hAnsi"), font_name_latin)


def set_document_fonts(doc, size=10):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "List Bullet", "List Number", "Quote"]:
        try:
            style = doc.styles[style_name]
            font = style.font
            font.name = LATIN_FONT
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
            rFonts.set(qn("w:ascii"), LATIN_FONT)
            rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        except KeyError:
            pass

    normal_font = doc.styles["Normal"].font
    normal_font.size = Pt(size)


INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE = re.compile(r"`([^`]+?)`")
INLINE_ITALIC = re.compile(r"(?<!\*)\*([^\*\n]+?)\*(?!\*)")


def add_inline_runs(paragraph, text, bold_override=False):
    """Parse inline markdown (bold, code, italic) and emit runs.
    Simple state-machine tokenizer: **bold**, `code`, *italic*.
    """
    pos = 0
    # Build token list
    tokens = []
    idx = 0
    while idx < len(text):
        m_bold = INLINE_BOLD.search(text, idx)
        m_code = INLINE_CODE.search(text, idx)
        m_ital = INLINE_ITALIC.search(text, idx)
        candidates = [(m.start(), m, kind) for m, kind in [(m_bold, "bold"), (m_code, "code"), (m_ital, "italic")] if m]
        if not candidates:
            tokens.append(("plain", text[idx:]))
            break
        candidates.sort(key=lambda t: t[0])
        start, m, kind = candidates[0]
        if start > idx:
            tokens.append(("plain", text[idx:start]))
        tokens.append((kind, m.group(1)))
        idx = m.end()

    for kind, content in tokens:
        run = paragraph.add_run(content)
        if bold_override:
            run.bold = True
        if kind == "bold":
            run.bold = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "Consolas")
            rFonts.set(qn("w:ascii"), "Consolas")
            rFonts.set(qn("w:hAnsi"), "Consolas")
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        elif kind == "italic":
            run.italic = True


def strip_frontmatter(lines):
    if not lines or lines[0].strip() != "---":
        return lines, []
    out = []
    fm = []
    in_fm = True
    for i, line in enumerate(lines[1:], start=1):
        if in_fm and line.strip() == "---":
            return lines[i + 1:], fm
        if in_fm:
            fm.append(line)
    return lines, []


def parse_table_row(line):
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]


def is_table_separator(line):
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = parse_table_row(stripped)
    return all(re.match(r"^:?-{3,}:?$", c) for c in cells if c)


def convert(md_path: Path, docx_path: Path, title_hint: str = None):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    lines, _ = strip_frontmatter(lines)

    doc = Document()
    set_document_fonts(doc, size=10)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    if title_hint:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_hint)
        run.bold = True
        run.font.size = Pt(14)

    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                in_code_block = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.append(rFonts)
                rFonts.set(qn("w:eastAsia"), "Consolas")
                rFonts.set(qn("w:ascii"), "Consolas")
                rFonts.set(qn("w:hAnsi"), "Consolas")
                p.paragraph_format.left_indent = Inches(0.2)
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Horizontal rule
        if stripped in ("---", "***", "___") and i > 0:
            doc.add_paragraph().add_run("─" * 60)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            content = m.group(2).strip()
            if level > 4:
                level = 4
            h = doc.add_heading("", level=level)
            add_inline_runs(h, content, bold_override=True)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and "|" in stripped[1:]:
            # Look ahead: next line should be separator
            if i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                header_cells = parse_table_row(stripped)
                i += 2
                data_rows = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    data_rows.append(parse_table_row(lines[i].strip()))
                    i += 1
                table = doc.add_table(rows=1 + len(data_rows), cols=len(header_cells))
                table.style = "Light Grid Accent 1"
                for c, h_text in enumerate(header_cells):
                    cell = table.rows[0].cells[c]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline_runs(p, h_text, bold_override=True)
                    set_cell_font(cell, bold=True, size=9)
                for r, row in enumerate(data_rows, start=1):
                    for c, cell_text in enumerate(row):
                        if c >= len(table.rows[r].cells):
                            break
                        cell = table.rows[r].cells[c]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_inline_runs(p, cell_text)
                        set_cell_font(cell, size=9)
                doc.add_paragraph()  # spacer
                continue

        # Block quote
        if stripped.startswith(">"):
            quote_text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline_runs(p, quote_text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # Bullet list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            content = m.group(2)
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
                p.add_run("• ")
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.5 + 0.25 * (indent // 2))
            add_inline_runs(p, content)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            content = m.group(3)
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
                p.add_run(f"{m.group(2)}. ")
            add_inline_runs(p, content)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")


def main():
    parser = argparse.ArgumentParser(description="Markdown → DOCX converter")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--title", default=None)
    args = parser.parse_args()

    if not args.input.exists():
        print(f"ERROR: {args.input} not found", file=sys.stderr)
        sys.exit(1)

    convert(args.input, args.output, title_hint=args.title)


if __name__ == "__main__":
    main()
